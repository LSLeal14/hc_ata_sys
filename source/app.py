import streamlit as st
from docx import Document
import re
import streamlit.components.v1 as components

# Extrai os campos {{campo}} do .docx
def extrair_campos(doc):
    texto = "\n".join([p.text for p in doc.paragraphs])
    return list(set(re.findall(r"\{\{(.*?)\}\}", texto)))

# Converte o documento para HTML simples, substituindo os campos
def docx_para_html(doc, dados):
    html = ""
    for p in doc.paragraphs:
        linha = p.text
        for chave, valor in dados.items():
            linha = linha.replace(f"{{{{{chave}}}}}", valor)
        html += f"<p>{linha}</p>\n"
    return f"""
    <html>
    <head>
        <meta charset="utf-8">
    </head>
    <body>
        {html}
    </body>
    </html>
    """

st.set_page_config(page_title="Word para PDF no Navegador", layout="centered")
st.title("📝 Preencher Template Word e Gerar PDF")

uploaded_file = st.file_uploader("Envie o template .docx com campos {{nome}}, {{data}}, etc.", type="docx")

if uploaded_file:
    doc = Document(uploaded_file)
    campos = extrair_campos(doc)

    if campos:
        st.success(f"Campos encontrados: {', '.join(campos)}")
        dados = {}
        for campo in campos:
            dados[campo] = st.text_input(campo)

        if st.button("Gerar PDF"):
            html_resultado = docx_para_html(doc, dados)

            # Passa o HTML para o componente React (frontend)
            components.html(f"""
            <html>
            <head>
            <script src="https://cdnjs.cloudflare.com/ajax/libs/jspdf/2.5.1/jspdf.umd.min.js"></script>
            </head>
            <body>
            <h3>Visualização do documento:</h3>
            <div id="content" style="border:1px solid #ddd;padding:10px;margin-bottom:10px;">{html_resultado}</div>
            <button id="download">📥 Baixar PDF</button>

            <script>
            const {{ jsPDF }} = window.jspdf;

            document.getElementById("download").onclick = () => {{
                var doc = new jsPDF();
                var content = document.getElementById("content").innerText;
                doc.text(content, 10, 10);
                doc.save("documento_preenchido.pdf");
            }};
            </script>
            </body>
            </html>
            """, height=400)

    else:
        st.warning("Nenhum campo {{campo}} encontrado no documento.")
